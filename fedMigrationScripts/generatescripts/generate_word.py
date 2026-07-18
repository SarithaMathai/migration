#!/usr/bin/env python3
"""
Generate FederatedGqlBreakDown-{domain}.docx — Microsoft Word format.

Styles match the reference docs in 'final PO BreakDown Doc/':
  - Title: bold 20pt blue
  - 4-column metrics banner (light-blue cells)
  - Section headings 14pt blue, sub-headings 12pt
  - Story tables with colored headers and complexity-coded text
  - Emoji icons throughout (paste into Confluence cleanly)

The per-domain doc is ONE merged Backend + Frontend page (mirrors generate_breakdown.py's
.md merge): '## Backend' (build_word_doc's existing content) then '## Frontend'
(generate_frontend.py's build_fe_section(), rendered via render_md_block).

Output:  output/summary/{domain}/FederatedGqlBreakDown-{domain}.docx  (per-domain subfolder)
         output/summary/Federated+Graphql+Stories+-+BreakDown.docx    (global, summary/ root)

Run:
    python generate_word.py              # all phase-1 domains + global
    python generate_word.py attachment   # single domain
    python generate_word.py --global     # global doc only
"""

import re
import sys
import importlib.util
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Load shared parsers from generate_breakdown ──────────────────────────────
HERE = Path(__file__).resolve().parent
OUTPUT = HERE.parent.parent / "output" / "summary"   # generated docs go to migration/output/summary/

def _load(name: str):
    spec = importlib.util.spec_from_file_location(name, HERE / f"{name}.py")
    mod  = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod

bd = _load("generate_breakdown")

ALL_DOMAINS    = bd.ALL_DOMAINS
DOMAIN_LABELS  = bd.DOMAIN_LABELS
DGS_MAP        = bd.DGS_MAP
PHASE_ICONS    = bd.PHASE_ICONS
PHASE_NAMES    = bd.PHASE_NAMES
TYPE_ICONS     = bd.TYPE_ICONS
COMPLEXITY_ICONS = bd.COMPLEXITY_ICONS
SIZE_MAP       = bd.SIZE_MAP

# ─── Color palette (matching reference Word docs) ─────────────────────────────
C_TITLE      = "1F497D"   # title + section headings (medium blue)
C_BODY       = "373737"   # body text
C_SUBTLE     = "444444"   # subtitle italic
C_BANNER_BG  = "DAE8FC"   # 4-column metrics banner background
C_BANNER_LBL = "444477"   # banner label text
C_TBL_BORDER = "BFBFBF"   # table border color
C_TBL_HDR_BG = "DAE8FC"   # story-table header background (light blue — icons/text stay legible)
C_TBL_HDR_FG = "1F3864"   # header text dark navy (dark-on-light)
C_ALT_ROW    = "F5F8FF"   # subtle alternate row tint

# Phase header cell backgrounds
PHASE_BG = {
    "B": "D6E4F0", "C": "D6EDDA", "D": "FFF3CD",
    "E": "FFE0CC", "F": "EDE0F8", "G": "D4EFDF",
}
PHASE_FG = {
    "B": "1F3864", "C": "1A4731", "D": "7D5002",
    "E": "7A2104", "F": "4A0F7A", "G": "0A4023",
}

# Complexity text colors
CL_COLOR = {
    "very high": "C0392B",
    "high":      "D35400",
    "medium":    "9B7A00",
    "low":       "1E7A34",
}


# ─── Word helpers ─────────────────────────────────────────────────────────────
def rgb(hex_color: str) -> RGBColor:
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


def _get_or_add(parent, tag: str):
    """Get existing child element by tag, or create and prepend it."""
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.insert(0, el)
    return el


def set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = _get_or_add(tc, "w:tcPr")
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_table_borders(table, color: str = C_TBL_BORDER, sz: str = "4") -> None:
    tbl   = table._tbl
    tblPr = _get_or_add(tbl, "w:tblPr")
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        borders.append(b)
    tblPr.append(borders)


def set_col_widths(table, widths_inches: list[float]) -> None:
    """Set column widths (in inches) for every column."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(widths_inches):
                break
            tc   = cell._tc
            tcPr = _get_or_add(tc, "w:tcPr")
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(int(widths_inches[i] * 1440)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def set_row_height(row, height_pt: float) -> None:
    tr   = row._tr
    trPr = _get_or_add(tr, "w:trPr")
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),  str(int(height_pt * 20)))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)


def cell_para(cell, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Return (creating if needed) the first paragraph of a cell."""
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.clear()
    p.alignment = alignment
    cell.vertical_alignment = 0   # top
    return p


def add_run(para, text: str, bold=False, italic=False,
            size_pt=None, color_hex=None, mono=False) -> None:
    run = para.add_run(text)
    if bold:      run.bold      = True
    if italic:    run.italic    = True
    if mono:      run.font.name = "Courier New"
    if size_pt:   run.font.size = Pt(size_pt)
    if color_hex: run.font.color.rgb = rgb(color_hex)


def inline_md(para, text: str, size_pt: float = 10,
              color_hex: str = C_BODY, bold_base=False) -> None:
    """Render inline Markdown (backticks → mono, **bold**, *italic*) as Word runs."""
    # Split on `code`, **bold**, *italic*
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(size_pt)
        elif part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(size_pt)
            r.font.color.rgb = rgb(color_hex)
        elif part.startswith("*") and part.endswith("*"):
            r = para.add_run(part[1:-1])
            r.italic = True
            r.font.size = Pt(size_pt)
            r.font.color.rgb = rgb(color_hex)
        else:
            r = para.add_run(part)
            if bold_base: r.bold = True
            r.font.size = Pt(size_pt)
            r.font.color.rgb = rgb(color_hex)


def parse_md_table(md_text: str) -> tuple[list[str], list[list[str]]]:
    """Parse Markdown table into (headers, data_rows)."""
    lines = [l.strip() for l in md_text.split("\n") if l.strip().startswith("|")]
    if not lines:
        return [], []
    def row_cells(line: str) -> list[str]:
        return [c.strip() for c in line.split("|")[1:-1]]
    headers = row_cells(lines[0])
    rows    = [row_cells(l) for l in lines[2:] if not re.match(r"^\|[-:\s|]+\|$", l)]
    return headers, [r for r in rows if any(c.strip() for c in r)]


# ─── Section writers ──────────────────────────────────────────────────────────
def section_heading(doc: Document, text: str, level: int = 2) -> None:
    """Add a section heading. level 1=20pt, 2=14pt, 3=12pt."""
    sizes  = {1: 20, 2: 14, 3: 12}
    colors = {1: C_TITLE, 2: C_TITLE, 3: "373737"}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(sizes.get(level, 12))
    r.font.color.rgb = rgb(colors.get(level, C_TITLE))


def add_plain_table(doc: Document, headers: list[str], rows: list[list[str]],
                    col_widths: list[float] | None = None) -> None:
    """Add a generic table with colored header row and clean borders."""
    if not headers and not rows:
        return
    ncols = max(len(headers), max((len(r) for r in rows), default=0))
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    set_table_borders(table)
    table.style = doc.styles["Table Grid"]

    # Header row
    hdr_row = table.rows[0]
    set_row_height(hdr_row, 18)
    for i, h in enumerate(headers):
        if i >= ncols: break
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_TBL_HDR_BG)
        p = cell_para(cell)
        add_run(p, h, bold=True, size_pt=9, color_hex=C_TBL_HDR_FG)

    # Data rows
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        if ri % 2 == 1:
            for cell in trow.cells:
                set_cell_bg(cell, C_ALT_ROW)
        for ci, val in enumerate(row):
            if ci >= ncols: break
            p = cell_para(trow.cells[ci])
            inline_md(p, val or "—", size_pt=9)

    if col_widths:
        set_col_widths(table, col_widths)
    doc.add_paragraph()   # spacer


def add_md_section(doc: Document, heading: str, md_text: str,
                   col_widths: list[float] | None = None) -> None:
    """Extract a Markdown table from md_text and add it as a Word table."""
    if not md_text:
        return

    # Find first table (lines starting with |)
    table_lines = [l for l in md_text.split("\n") if l.strip().startswith("|")]
    prose_lines  = [l for l in md_text.split("\n") if l.strip() and not l.strip().startswith("|")
                    and not re.match(r"^[-=*]{3,}", l.strip())]

    section_heading(doc, heading)

    # Prose paragraphs first (e.g., notes below tables)
    for line in prose_lines:
        clean = re.sub(r"^>\s*", "", line).strip()
        if clean:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            inline_md(p, clean, size_pt=9, color_hex=C_SUBTLE, bold_base=False)

    # Table
    if table_lines:
        headers, rows = parse_md_table("\n".join(table_lines))
        add_plain_table(doc, headers, rows, col_widths)


def render_md_block(doc: Document, md_lines: list[str]) -> None:
    """Render a small Markdown block (## headings, paragraphs, - bullets, > quotes, | tables |)
    into the doc — used so the global breakdown's Overview/Glossary/Phases/T-Shirt matches the .md."""
    i, n = 0, len(md_lines)
    while i < n:
        s = md_lines[i].strip()
        if s.startswith("### "):
            section_heading(doc, s[4:].strip(), level=3)
        elif s.startswith("## "):
            section_heading(doc, s[3:].strip())
        elif s.startswith("|"):
            tbl = []
            while i < n and md_lines[i].strip().startswith("|"):
                tbl.append(md_lines[i]); i += 1
            headers, rows = parse_md_table("\n".join(tbl))
            add_plain_table(doc, headers, rows)
            continue
        elif s.startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
            inline_md(p, s[2:].strip(), size_pt=9, color_hex=C_SUBTLE)
        elif s.startswith("- "):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
            add_run(p, "•  ", size_pt=9, color_hex=C_SUBTLE); inline_md(p, s[2:].strip(), size_pt=9)
        elif s and s != "---":
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
            inline_md(p, s, size_pt=9.5)
        i += 1


# ─── Metrics banner (4-col, light-blue cells) ─────────────────────────────────
def add_metrics_banner(doc: Document, domain: str, stories: list) -> None:
    label    = DOMAIN_LABELS[domain]
    dgs      = DGS_MAP[domain]
    ts       = bd.tshirt(domain)
    total    = len(stories)
    vh       = sum(1 for s in stories if s["complexity"] == "very high")
    hi       = sum(1 for s in stories if s["complexity"] == "high")
    me       = sum(1 for s in stories if s["complexity"] == "medium")
    lo       = sum(1 for s in stories if s["complexity"] == "low")
    by_phase = bd.group_by_phase(stories)
    phases   = ", ".join(k for k in sorted(by_phase.keys()))

    # 4-col banner table
    cells_data = [
        ("Total Stories", str(total)),
        ("T-Shirt Size",  ts),
        ("Complexity",    f"🔴 {vh}  🟠 {hi}  🟡 {me}  🟢 {lo}"),
        ("Phases",        phases),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = doc.styles["Table Grid"]
    set_table_borders(table, color="A0B4CF", sz="2")
    row   = table.rows[0]
    set_row_height(row, 32)

    for i, (lbl, val) in enumerate(cells_data):
        cell = row.cells[i]
        set_cell_bg(cell, C_BANNER_BG)
        p = cell_para(cell, WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p, lbl + "\n", size_pt=8, color_hex=C_BANNER_LBL)
        add_run(p, val, bold=True, size_pt=11, color_hex=C_TITLE)

    doc.add_paragraph()   # spacer

    # Sub-header info row
    info_p = doc.add_paragraph()
    add_run(info_p, f"Target DGS: ", bold=True, size_pt=9, color_hex=C_BODY)
    add_run(info_p, dgs, mono=True,  size_pt=9, color_hex="1F497D")
    add_run(info_p, f"   |   Generated: {date.today().isoformat()}", size_pt=9, color_hex=C_SUBTLE)
    info_p.paragraph_format.space_after = Pt(4)

    # Icon legend
    legend_p = doc.add_paragraph()
    add_run(legend_p,
            "🔷 Query · 🔶 Mutation · 🔸 Field Resolver   "
            "🔴 Very High · 🟠 High · 🟡 Medium · 🟢 Low   "
            "🔴🔬 spike-gated · 📖 B · 🔍 C · ✏️ D · ⚙️ E · 🔗 F · 🧪 G",
            size_pt=8, color_hex=C_SUBTLE, italic=True)
    legend_p.paragraph_format.space_after = Pt(8)


# ─── Story table (one row per story — Intent/Today/Done-when in the AC column) ─
def _first_para(cell):
    """Return the cell's first (empty) paragraph, ready to receive runs."""
    p = cell.paragraphs[0]
    p.clear()
    p.paragraph_format.space_after = Pt(1)
    return p


def _fill_ac_cell(cell, s) -> None:
    """Write Intent, Today and bulleted Done-when into a single AC cell as stacked paragraphs."""
    first = True
    def para(indent=0.0):
        nonlocal first
        p = _first_para(cell) if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(1)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        return p

    wrote = False
    if s.get("intent"):
        p = para()
        add_run(p, "Intent — ", bold=True, size_pt=8.5, color_hex=C_BODY)
        inline_md(p, s["intent"], size_pt=8.5)
        wrote = True
    today = bd.minimal_logic(s)
    if today and today != "—":
        p = para()
        add_run(p, "Today — ", bold=True, size_pt=8.5, color_hex=C_BODY)
        inline_md(p, today, size_pt=8.5)
        wrote = True
    if s["ac"]:
        p = para()
        add_run(p, "Done when:", bold=True, size_pt=8.5, color_hex=C_BODY)
        for a in s["ac"]:
            bp = para(indent=0.12)
            add_run(bp, "• ", size_pt=8.5, color_hex=C_SUBTLE)
            inline_md(bp, a, size_pt=8.5)
        wrote = True
    if not wrote:
        add_run(_first_para(cell), "—", size_pt=8.5, color_hex=C_SUBTLE)


def _fill_tests_cell(cell, s) -> None:
    """Write the ☐ Key-Test checklist for a High/VH story; '—' for the rest."""
    if s["tests"] and s["complexity"] in ("high", "very high"):
        first = True
        for t in s["tests"]:
            p = _first_para(cell) if first else cell.add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(1)
            add_run(p, "☐  ", size_pt=8.5)
            inline_md(p, t, size_pt=8.5)
    else:
        add_run(_first_para(cell), "—", size_pt=8.5, color_hex=C_SUBTLE)


def add_story_table(doc: Document, phase_key: str, stories: list) -> None:
    """Phase heading band + a table with one row per story. Intent, Today and Done-when share
    the **Acceptance Criteria** column; **Key Tests** is a separate column present ONLY for
    complex phases (those with a High / Very High story)."""
    if not stories:
        return

    phase_icon = PHASE_ICONS.get(phase_key, "📄")
    phase_name = PHASE_NAMES.get(phase_key, f"Phase {phase_key}")
    bg         = PHASE_BG.get(phase_key, "E8E8E8")
    fg         = PHASE_FG.get(phase_key, "333333")

    # Phase heading (shaded band)
    ph_p = doc.add_paragraph()
    ph_p.paragraph_format.space_before = Pt(10)
    ph_p.paragraph_format.space_after  = Pt(3)
    add_run(ph_p, f"{phase_icon} Phase {phase_key} — {phase_name}  ({len(stories)} stories)",
            bold=True, size_pt=11, color_hex=fg)
    pPr = _get_or_add(ph_p._p, "w:pPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), bg)
    pPr.append(shd)

    # A "complex phase" carries the Key Tests column; simple phases omit it.
    complex_phase = any(s["complexity"] in ("high", "very high") for s in stories)
    headers = ["Story", "Complexity", "Type / Calls", "Depends On", "Acceptance Criteria"]
    if complex_phase:
        headers.append("Key Tests")
    ncols = len(headers)

    table = doc.add_table(rows=1 + len(stories), cols=ncols)
    table.style = doc.styles["Table Grid"]
    set_table_borders(table)

    # Header row
    hdr = table.rows[0]
    set_row_height(hdr, 16)
    for i, h in enumerate(headers):
        set_cell_bg(hdr.cells[i], C_TBL_HDR_BG)
        add_run(_first_para(hdr.cells[i]), h, bold=True, size_pt=9, color_hex=C_TBL_HDR_FG)

    # Data rows
    for ri, s in enumerate(stories, 1):
        cl     = s["complexity"]
        cicon  = COMPLEXITY_ICONS.get(cl, "⚪")
        size   = SIZE_MAP.get(cl, "?")
        ticon  = TYPE_ICONS.get(s["type"], "📄")
        tname  = s["type"].replace("field resolver", "Field Resolver").title()
        cl_col = CL_COLOR.get(cl, C_BODY)
        spk    = bd.spike_for(s)
        cells  = table.rows[ri].cells

        if ri % 2 == 0:                            # subtle zebra striping
            for c in cells:
                set_cell_bg(c, C_ALT_ROW)

        # Story cell — icon(s) + id + title, plus a spike-gated note.
        sp = _first_para(cells[0])
        if spk:
            add_run(sp, "🔴🔬 ", size_pt=9)
        add_run(sp, f"{ticon} ", size_pt=9)
        add_run(sp, s["id"], bold=True, mono=True, size_pt=8.5,
                color_hex=(CL_COLOR["very high"] if spk else C_TITLE))
        tp = cells[0].add_paragraph()
        tp.paragraph_format.space_after = Pt(1)
        inline_md(tp, s["title"], size_pt=8.5)
        if spk:
            cp = cells[0].add_paragraph()
            cp.paragraph_format.space_after = Pt(1)
            add_run(cp, f"🔴🔬 Spike-gated on SPIKE-{spk} ({bd.SPIKE_TITLES.get(spk, '')})",
                    italic=True, size_pt=8, color_hex=CL_COLOR["very high"])

        # Complexity
        add_run(_first_para(cells[1]), f"{cicon} {cl.title()} [{size}]", size_pt=8.5, color_hex=cl_col)

        # Type / Calls
        cp = _first_para(cells[2])
        add_run(cp, tname, size_pt=8.5, color_hex=C_SUBTLE)
        if s["ext_services"]:
            cc = cells[2].add_paragraph()
            cc.paragraph_format.space_after = Pt(1)
            add_run(cc, "Calls: ", size_pt=8, color_hex=C_SUBTLE)
            add_run(cc, ", ".join(s["ext_services"]), size_pt=8, color_hex=C_BODY)

        # Depends On
        dep = s["depends"] or "—"
        if spk:
            sref = f"SPIKE-{spk}"
            dep  = sref if dep in ("—", "") else f"{sref}, {dep}"
        add_run(_first_para(cells[3]), dep, size_pt=8.5, color_hex=C_SUBTLE)

        # Acceptance Criteria (Intent / Today / Done when)
        _fill_ac_cell(cells[4], s)

        # Key Tests (complex phases only)
        if complex_phase:
            _fill_tests_cell(cells[5], s)

    if complex_phase:
        set_col_widths(table, [1.6, 0.75, 1.15, 0.9, 3.0, 1.7])
    else:
        set_col_widths(table, [1.7, 0.8, 1.3, 0.95, 4.15])
    doc.add_paragraph()

    # DGS init notes (B-01-type) — callouts beneath the table.
    note_items = [s for s in stories if s["note"] and bd.strip_noise(s["note"]).strip()]
    for s in note_items:
        nc = bd.strip_noise(s["note"]).lstrip("> ").strip()
        np = doc.add_paragraph()
        np.paragraph_format.left_indent = Inches(0.15)
        np.paragraph_format.space_after = Pt(2)
        add_run(np, "ℹ️  ", size_pt=9)
        inline_md(np, f"{s['id']}: {nc}", size_pt=8.5, color_hex=C_SUBTLE)
    if note_items:
        doc.add_paragraph()


# ─── Domain spike map (relates this domain's stories to the program spikes) ───
def add_domain_spike_map(doc: Document, stories: list) -> None:
    """Compact 'gated story → program spike' map. The full brief/steps/resolver detail
    live once in the global doc; this only links to it (no duplicated context)."""
    gated = [(s, bd.spike_for(s)) for s in stories]
    gated = [(s, b) for s, b in gated if b]

    section_heading(doc, "Spikes & Complex Cases")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_run(p,
            "This domain's complex, cross-cutting problems are tracked once as program spikes in the global "
            "breakdown — see 'Phase 0 — Program Spikes' and 'Spike Detail' (the brief, the decision, "
            "intended cross-domain steps, and every affected resolver's external deps + current logic) in "
            "Federated+Graphql+Stories+-+BreakDown. Nothing from there is repeated here; the stories below just "
            "link to it. Follow a story's SPIKE-0x id to the global Spike Detail for the target flow + the "
            "external services each resolver calls (see 'How to read the spikes & related stories' in the global doc).",
            size_pt=9, color_hex=C_SUBTLE, italic=True)

    if gated:
        add_plain_table(
            doc, ["Story", "Program spike", "Bucket"],
            [[f"🔴🔬 {s['id']} — {re.sub('`', '', s['title'])}", f"SPIKE-{b}", bd.SPIKE_TITLES[b]]
             for s, b in gated],
            col_widths=[4.6, 1.5, 2.4],
        )
    else:
        np = doc.add_paragraph()
        add_run(np, "No spike-gated stories in this domain.", size_pt=9, italic=True, color_hex=C_SUBTLE)

    np2 = doc.add_paragraph()
    np2.paragraph_format.space_after = Pt(6)
    add_run(np2,
            "Simple, intuitive decisions (drift-op cleanup, dead-method audits, auth-token parity, sort/DTO "
            "shape) are resolved inline in the owning story — they are not spikes.",
            size_pt=9, italic=True, color_hex=C_SUBTLE)


# ─── Domain Word doc builder ──────────────────────────────────────────────────
def build_word_doc(domain: str) -> Document:
    label   = DOMAIN_LABELS[domain]
    src_dir = bd.get_domain_dir(domain)
    # Drop Phase-S spikes — centralized as program spikes (see global doc).
    stories = [s for s in bd.parse_stories(src_dir / "be-04-stories.md") if s["phase"] != "S"]
    po      = bd.read_po_sections(src_dir / "be-04-po-summary.md")
    by_phase = bd.group_by_phase(stories)

    doc = Document()

    # Page margins (narrower to fit the story table)
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(1.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Title ──────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    add_run(title_p, f"Federated GraphQL Breakdown — {label}",
            bold=True, size_pt=20, color_hex=C_TITLE)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(8)
    add_run(sub_p,
            f"Target DGS: {DGS_MAP[domain]}  ·  "
            "spark-internal-graphql → Netflix DGS Federation",
            italic=True, size_pt=11, color_hex=C_SUBTLE)

    # ── Metrics banner ─────────────────────────────────────────────────────
    add_metrics_banner(doc, domain, stories)

    # ── Backend section marker (mirrors the '## Backend' heading in the merged .md) ──
    section_heading(doc, "Backend", level=1)

    # ── §1 What Are We Building? ───────────────────────────────────────────
    if po.get("what"):
        section_heading(doc, "What Are We Building?")
        clean = bd.strip_noise(po["what"])
        for para_block in clean.split("\n\n"):
            para_block = para_block.strip()
            if not para_block:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            inline_md(p, para_block.replace("\n", " "), size_pt=10)

    # ── §2 Migration Scope ─────────────────────────────────────────────────
    if po.get("scope"):
        add_md_section(doc, "Migration Scope", po["scope"],
                       col_widths=[2.2, 0.7, 5.8])

    # ── Spikes & Complex Cases — map to program spikes (no duplicated detail) ─
    add_domain_spike_map(doc, stories)

    # ── §3 Effort Snapshot ─────────────────────────────────────────────────
    if po.get("phases") or po.get("capacity"):
        section_heading(doc, "Effort Snapshot by Phase")
        if po.get("phases"):
            hdrs, rows = parse_md_table(po["phases"])
            if hdrs:
                add_plain_table(doc, hdrs, rows, col_widths=[0.6, 3.0, 0.7, 2.5])

        # Capacity planning
        for line in (po.get("phases") or "").split("\n"):
            line = line.strip()
            if line.startswith(">") and "sprint" in line.lower():
                note_text = re.sub(r"^>\s*", "", line)
                np = doc.add_paragraph()
                add_run(np, "ℹ️  ", size_pt=9)
                inline_md(np, note_text, size_pt=9, color_hex=C_SUBTLE)

        if po.get("capacity"):
            section_heading(doc, "Capacity Planning", level=3)
            hdrs, rows = parse_md_table(po["capacity"])
            if hdrs:
                add_plain_table(doc, hdrs, rows, col_widths=[2.0, 2.0, 4.0])

    # ── §4 Sprint Sequencing ───────────────────────────────────────────────
    if po.get("sprints"):
        add_md_section(doc, "Recommended Sprint Sequencing", po["sprints"],
                       col_widths=[0.9, 1.5, 5.5])

    # ── §4b Recommended Implementation Order (same builder as the .md) ─────
    render_md_block(doc, bd.implementation_order_md(stories))

    # ── §4c Recommended Story Graph — 2 Backend Engineers (same builder) ───
    render_md_block(doc, bd.team_plan_md(stories))

    # ── §5 Stories by Phase ────────────────────────────────────────────────
    section_heading(doc, "Jira Stories by Phase")
    note_p = doc.add_paragraph()
    add_run(note_p,
            "Each row is one Jira story. Complexity drives T-shirt sizing in refinement. "
            "'Depends On' lists blocking story IDs within this domain.",
            size_pt=9, color_hex=C_SUBTLE, italic=True)
    note_p.paragraph_format.space_after = Pt(6)

    if not stories:
        doc.add_paragraph("No stories found.")
    else:
        for phase_key in sorted(by_phase.keys()):
            if phase_key == "S":
                continue   # spikes rendered in Phase 0 above
            add_story_table(doc, phase_key, by_phase[phase_key])

    # Decisions Required / Complex Story Breakdowns intentionally REMOVED — complex cases are
    # the program spikes (global doc), simple decisions fold into their owning story's AC.

    return doc


# ─── Global Word doc builder ──────────────────────────────────────────────────
def build_global_word(domains: "list[str] | None" = None, scope_label: str = "All Domains") -> Document:
    today = date.today().isoformat()
    domains = domains or ALL_DOMAINS
    doc   = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    tp = doc.add_paragraph()
    add_run(tp, f"Federated GraphQL — Migration Overview ({scope_label})",
            bold=True, size_pt=20, color_hex=C_TITLE)

    sp = doc.add_paragraph()
    add_run(sp,
            f"Generated: {today}  ·  {len(domains)} domains  ·  "
            "spark-internal-graphql → Netflix DGS via Hive Schema Registry",
            italic=True, size_pt=11, color_hex=C_SUBTLE)
    sp.paragraph_format.space_after = Pt(10)

    # Program cover banner (same shape as the per-domain metrics banner)
    _n_dgs = len({DGS_MAP[d].split(" (")[0] for d in domains})
    _g = [0, 0, 0, 0, 0]
    for _d in domains:
        try:
            _st = [s for s in bd.parse_stories(bd.get_domain_dir(_d) / "be-04-stories.md") if s["phase"] != "S"]
        except Exception:
            _st = []
        _g[0] += len(_st)
        _g[1] += sum(1 for s in _st if s["complexity"] == "very high")
        _g[2] += sum(1 for s in _st if s["complexity"] == "high")
        _g[3] += sum(1 for s in _st if s["complexity"] == "medium")
        _g[4] += sum(1 for s in _st if s["complexity"] == "low")
    _fe = bd.fe_story_stats()
    _fe_n  = sum(_fe.get(d, (0, 0, 0))[0] for d in domains)
    _fe_lo = sum(_fe.get(d, (0, 0, 0))[1] for d in domains)
    _fe_hi = sum(_fe.get(d, (0, 0, 0))[2] for d in domains)
    add_plain_table(doc, ["Program", "spark-internal-graphql → Netflix DGS Federation (Hive Registry)"], [
        ["Domains", str(len(domains))],
        ["Target DGS services", str(_n_dgs)],
        ["Total Backend Stories", str(_g[0])],
        ["Total Frontend Stories", f"{_fe_n} · {_fe_lo}–{_fe_hi}d single-engineer (Frontend section of each per-domain FederatedGqlBreakDown-<domain> page)"],
        ["Complexity (backend)", f"🔴 {_g[1]} Very High · 🟠 {_g[2]} High · 🟡 {_g[3]} Medium · 🟢 {_g[4]} Low"],
        ["Generated", today],
    ], col_widths=[2.2, 5.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Overview / Glossary / Phases / T-Shirt — same content as the .md
    render_md_block(doc, bd.program_overview_preamble())

    # Domain index table
    section_heading(doc, "Domain Index")

    idx_headers = ["#", "Domain", "Target DGS", "T-Shirt", "BE Stories", "🔴 VH", "🟠 H", "🟡 M", "🟢 L",
                   "FE Stories", "FE effort", "Breakdown pages"]
    idx_rows: list[list[str]] = []
    domain_story_map: dict[str, list] = {}
    grand = [0, 0, 0, 0, 0]  # total, vh, hi, me, lo

    for i, domain in enumerate(domains, 1):
        label = DOMAIN_LABELS[domain]
        ts    = bd.tshirt(domain)
        try:
            src_dir = bd.get_domain_dir(domain)
            stories = [s for s in bd.parse_stories(src_dir / "be-04-stories.md") if s["phase"] != "S"]
        except FileNotFoundError:
            stories = []
        domain_story_map[domain] = stories
        total = len(stories)
        vh    = sum(1 for s in stories if s["complexity"] == "very high")
        hi    = sum(1 for s in stories if s["complexity"] == "high")
        me    = sum(1 for s in stories if s["complexity"] == "medium")
        lo    = sum(1 for s in stories if s["complexity"] == "low")
        grand[0] += total; grand[1] += vh; grand[2] += hi; grand[3] += me; grand[4] += lo
        fc, flo, fhi = _fe.get(domain, (0, 0, 0))
        idx_rows.append([str(i), label, DGS_MAP[domain], ts,
                         str(total), str(vh), str(hi), str(me), str(lo),
                         str(fc), f"{flo}–{fhi}d",
                         f"FederatedGqlBreakDown-{domain}"])
    idx_rows.append(["", "TOTAL", "—", "—",
                     str(grand[0]), str(grand[1]), str(grand[2]), str(grand[3]), str(grand[4]),
                     str(_fe_n), f"{_fe_lo}–{_fe_hi}d", "—"])

    add_plain_table(doc, idx_headers, idx_rows,
                    col_widths=[0.3, 1.2, 1.5, 0.55, 0.6, 0.4, 0.4, 0.4, 0.4, 0.6, 0.7, 2.0])

    # Icon legend
    leg_p = doc.add_paragraph()
    add_run(leg_p,
            "🔷 Query · 🔶 Mutation · 🔸 Field Resolver   "
            "🔴 Very High · 🟠 High · 🟡 Medium · 🟢 Low   "
            "🔴🔬 spike-gated · 📖 B · 🔍 C · ✏️ D · ⚙️ E · 🔗 F · 🧪 G",
            size_pt=8, color_hex=C_SUBTLE, italic=True)
    leg_p.paragraph_format.space_after = Pt(10)

    # Spike Detail — plain-English brief + intended steps + per-resolver (external deps + current
    # logic). Rendered from the SAME builder as the .md so both stay identical.
    dd = [(d,) + (None,) * 8 + (domain_story_map.get(d, []),) for d in domains]
    render_md_block(doc, bd.build_spike_detail(dd))

    # Per-domain story detail intentionally NOT included — this is an overview.
    # Each domain's phase tables are in its own FederatedGqlBreakDown-<domain> breakdown.

    return doc


# ─── Frontend section (merged into the same per-domain .docx) ─────────────────
def _fe_section_lines_for(domain: str) -> "list[str] | None":
    """Load generate_frontend.py and build this domain's Frontend section markdown lines
    (same data as generate_breakdown.py's build_fe_section_for(), used for the .md merge),
    or None if fe-08-frontend-stories.md has no stories for this domain / doesn't exist."""
    try:
        fe_mod = _load("generate_frontend")
        stories = fe_mod.parse_fe_stories()
    except Exception:
        return None
    if not stories:
        return None
    from collections import defaultdict as _dd
    by_dom = _dd(list)
    for s in stories:
        by_dom[fe_mod.domain_key_from_token(s["id"].rsplit("-FE-", 1)[0])].append(s)
    group = sorted(by_dom.get(domain, []), key=lambda s: s["id"])
    if not group:
        return None
    ops = []
    try:
        client_defs = fe_mod.load_client_defs()
        usage_idx = fe_mod.load_usage_index()
        registry = fe_mod.build_schema_registry()
        be_stories = fe_mod.load_be_stories()
        ops, _frags, _coverage, _fragments = fe_mod.cross_reference(
            client_defs, registry, usage_idx, be_stories)
    except Exception:
        ops = []
    return fe_mod.build_fe_section(domain, group, ops)


# ─── Runner ────────────────────────────────────────────────────────────────────
def generate_word_for(domain: str) -> None:
    # Per-domain artifact lives in output/summary/{domain}/ — ONE merged .docx:
    # '## Backend' (build_word_doc) + '## Frontend' (generate_frontend.py's build_fe_section),
    # same merge convention as generate_breakdown.py's .md.
    domain_dir = OUTPUT / domain
    domain_dir.mkdir(parents=True, exist_ok=True)
    doc = build_word_doc(domain)

    fe_lines = _fe_section_lines_for(domain)
    if fe_lines:
        doc.add_page_break()
        section_heading(doc, "Frontend", level=1)
        render_md_block(doc, fe_lines)

    out_file = domain_dir / f"FederatedGqlBreakDown-{domain}.docx"
    doc.save(str(out_file))
    print(f"  OK {domain}/FederatedGqlBreakDown-{domain}.docx")


def generate_global_word() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc      = build_global_word()
    out_file = OUTPUT / "Federated+Graphql+Stories+-+BreakDown.docx"
    doc.save(str(out_file))
    print(f"  OK Federated+Graphql+Stories+-+BreakDown.docx")


# Scoped .docx for a hand-picked set of modules — the default global .docx is left untouched.
# Single source of truth: the list lives in generate_breakdown so the .md and .docx can never diverge.
from generate_breakdown import CUSTOM_DOMAINS


def generate_global_custom_word() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc      = build_global_word(CUSTOM_DOMAINS, scope_label="Selected Modules")
    out_file = OUTPUT / "Federated+Graphql+Stories+-+BreakDown_custom.docx"
    doc.save(str(out_file))
    print(f"  OK Federated+Graphql+Stories+-+BreakDown_custom.docx ({', '.join(CUSTOM_DOMAINS)})")


def main() -> None:
    args        = sys.argv[1:]
    global_only = "--global" in args
    custom_only = "--custom" in args
    targets     = [a for a in args if not a.startswith("--")]

    if custom_only:
        generate_global_custom_word()
        return
    if global_only:
        generate_global_word()
        return

    domains = targets if targets else ALL_DOMAINS
    today   = date.today().isoformat()
    print(f"\n=== Word doc generation — {today} ===\n")

    for domain in domains:
        if domain not in ALL_DOMAINS:
            print(f"  UNKNOWN '{domain}' — skipping")
            continue
        try:
            generate_word_for(domain)
        except Exception as e:
            import traceback
            print(f"  FAIL {domain}: {type(e).__name__}: {e}")
            traceback.print_exc()

    if not targets:
        try:
            generate_global_word()
        except Exception as e:
            print(f"  FAIL global: {type(e).__name__}: {e}")
        try:
            generate_global_custom_word()
        except Exception as e:
            print(f"  FAIL global_custom: {type(e).__name__}: {e}")

    print("\nDone.\n")


if __name__ == "__main__":
    main()
